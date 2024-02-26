#-----------------------
# Imports and dependencies
#-----------------------
from sys import argv
import logging
from docx import Document
from docx.table import _Cell as Cell_T
from docx.document import Document as Document_T
from pathlib import Path
from docx2pdf import convert
from pywintypes import com_error
import pyzipper
from pypdf import PdfWriter, PdfReader
from pypdf.errors import PyPdfError
from enum import Enum
from datetime import datetime

#-----------------------
# Enums 
#-----------------------
class ConversionType(Enum):
    ZIP = 0
    PDF = 1

#-----------------------
# Global variables
#-----------------------
DEFAULT_IN = Path(".") / "dokumenty"
OUT_PATH = Path(".") / "out"
LOG_LEVEL = logging.INFO
ENCRYPTION_ALGO = "AES-256"
CONVERSION_TYPE = ConversionType.PDF

#-----------------------
# Conversion step
#-----------------------
def encrypt_pdf(in_path: Path, out_path: Path, password: str):
    global ENCRYPTION_ALGO
    reader = PdfReader(in_path)
    writer = PdfWriter()

    for page in reader.pages:
        writer.add_page(page)
    
    writer.encrypt(password, algorithm=ENCRYPTION_ALGO)
    with open(str(out_path), "wb") as f:
        writer.write(f)
    logging.info(f'Created PDF "{out_path}", password: {password}')

def encrypt_zip(pdf_path: Path, zip_path: Path, password: str):
    with pyzipper.AESZipFile(
        zip_path, 
        'w',
        compression=pyzipper.ZIP_DEFLATED,
        encryption=pyzipper.WZ_AES) as z_out:
        z_out.setpassword(bytes(password, 'UTF-8'))
        z_out.write(pdf_path, pdf_path.name)
    logging.info(f'Created ZIP "{zip_path}", password: {password}')


def convert_encrypt(files: dict[str, Path]):
    global OUT_PATH, ZIP_CONVERSION
    for birth_num, doc_path in files.items():
        logging.debug(f'Birth num: {birth_num}, Path: "{doc_path}"')

        try:
            pdf_path = OUT_PATH / f"{doc_path.stem}.pdf"
            temp_path = OUT_PATH / f"{doc_path.stem}.temp.pdf"
            zip_path = OUT_PATH / f"{doc_path.stem}.zip"
            def cleanup():
                pdf_path.unlink(True)
                temp_path.unlink(True)
                zip_path.unlink(True)

            #-----------------
            # Driver switch
            #-----------------
            if CONVERSION_TYPE == ConversionType.ZIP:
                convert(str(doc_path), str(pdf_path))
                encrypt_zip(pdf_path, zip_path, birth_num)
                pdf_path.unlink()
            elif CONVERSION_TYPE == ConversionType.PDF:
                convert(str(doc_path), str(temp_path))
                encrypt_pdf(temp_path, pdf_path, birth_num)
                temp_path.unlink()
            cleanup = None

        except com_error:
            logging.error(f'Conversion error: "{doc_path}" failed converting to PDF')
        except pyzipper.BadZipFile:
            logging.error(f'Zipping error: "{doc_path}" failed creating ZIP')
        except PyPdfError:
            logging.error(f'PDF error: "{pdf_path}" failed encrypting PDF')
        except Exception as err:
            logging.error(f'Unknown error: "{doc_path}" {err!r}')

        finally:
            if cleanup is not None:
                cleanup()

#-----------------------
# Getting docs and data
#-----------------------
def get_birth_number(doc_path: Path):
    doc = Document(str(doc_path))
    if not isinstance(doc, Document_T):
        logging.error(f'Failed opening "{doc_path}"')
        return None
    if len(doc.tables) <= 2:
        logging.error(f'No tables in "{doc_path}"')
        return None
    cell = doc.tables[2].cell(1,2)
    if isinstance(cell, Cell_T):
        return cell.text.strip()
    logging.error(f'Cell not found in "{doc_path}"')
    return None

def birth_number_gen(doc_paths: list[Path]):
    for doc_path in doc_paths:
        birth_num = get_birth_number(doc_path)
        yield (doc_path, birth_num)

def get_docx_files(dir_name: Path):
    logging.info(f'{datetime.now():%d-%m-%Y} - Starting on folder "{dir_name}"')
    return list(dir_name.glob('**/*.docx'))

#-----------------------
# Setup functions
#-----------------------
def init_main():
    global LOG_LEVEL
    msg_format = "[%(levelname)s]: %(message)s"
    logging.basicConfig(level=LOG_LEVEL, filename="log.txt", encoding="UTF-8", format=msg_format)
    logging.getLogger().name = "Podilnici"

    OUT_PATH.mkdir(parents=True, exist_ok=True)
    if not OUT_PATH.is_dir():
        logging.error(f'"{OUT_PATH}" is not a directory')
    if len(argv) == 1:
        logging.debug(f'Adding "{DEFAULT_IN}" to command line arguments')
        argv.append(str(DEFAULT_IN))


#-----------------------
# Main function
#-----------------------
def main():
    init_main()

    files = get_docx_files(Path(argv[1]))
    birth_nums = {
        birth_num: file
        for file, birth_num
        in birth_number_gen(files)
        if birth_num is not None
    }

    convert_encrypt(birth_nums)

if __name__=='__main__':
    main()